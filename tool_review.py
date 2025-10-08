# tool_review.py - Refined Item Definition Reviewer Tools
import os
import json
from datetime import datetime
from cat.mad_hatter.decorators import tool
from cat.log import log

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log.warning("python-docx not available - DOCX reading disabled")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    log.warning("PyPDF2 not available - PDF reading disabled")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def load_checklist(plugin_folder):
    """
    Load ISO 26262 Part 3 compliance checklist from JSON.
    
    Args:
        plugin_folder (str): Path to plugin root
        
    Returns:
        dict: Checklist data
        
    Raises:
        FileNotFoundError: If checklist file not found
    """
    checklist_path = os.path.join(plugin_folder, "checklists", "item_definition_checklist.json")
    
    try:
        with open(checklist_path, "r", encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        log.error(f"Checklist not found: {checklist_path}")
        raise FileNotFoundError(f"Checklist file not found at {checklist_path}")
    except json.JSONDecodeError as e:
        log.error(f"Invalid checklist JSON: {e}")
        raise ValueError(f"Checklist file is corrupted: {e}")


def load_item_definition(plugin_folder):
    """
    Load Item Definition from file in item_definition_to_review folder.
    Supports .pdf, .docx, .txt formats.
    
    Args:
        plugin_folder (str): Path to plugin root
        
    Returns:
        str: Item definition content or None if not found
    """
    item_def_folder = os.path.join(plugin_folder, "item_definition_to_review")
    
    if not os.path.exists(item_def_folder):
        log.error(f"Folder not found: {item_def_folder}")
        return None
    
    supported_extensions = {'.pdf', '.docx', '.txt'}
    
    # Find first supported file
    for filename in os.listdir(item_def_folder):
        file_path = os.path.join(item_def_folder, filename)
        
        if not os.path.isfile(file_path):
            continue
        
        _, ext = os.path.splitext(filename.lower())
        
        if ext not in supported_extensions:
            continue
        
        log.info(f"📄 Found file: {filename}")
        
        # Extract content based on type
        content = extract_file_content(file_path, ext)
        
        if content and content.strip():
            log.info(f"✅ Extracted {len(content)} characters from {filename}")
            return content
        else:
            log.warning(f"⚠️ File {filename} appears empty or unreadable")
    
    log.warning("No valid Item Definition files found")
    return None


def extract_file_content(file_path, file_extension):
    """
    Extract text content from different file formats.
    
    Args:
        file_path (str): Path to file
        file_extension (str): Extension (.pdf, .docx, .txt)
        
    Returns:
        str: Extracted text or empty string on error
    """
    try:
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.pdf':
            if not PDF_AVAILABLE:
                log.error("PyPDF2 not installed - cannot read PDF")
                return ""
            
            content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            content.append(text)
                    except Exception as e:
                        log.warning(f"Failed to extract page {page_num}: {e}")
            
            return "\n".join(content)
        
        elif file_extension == '.docx':
            if not DOCX_AVAILABLE:
                log.error("python-docx not installed - cannot read DOCX")
                return ""
            
            doc = DocxDocument(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            return "\n".join(content)
        
        else:
            log.warning(f"Unsupported file extension: {file_extension}")
            return ""
    
    except Exception as e:
        log.error(f"Error extracting content from {file_path}: {e}")
        return ""


def format_checklist_for_llm(checklist):
    """
    Format checklist in a clear, structured way for LLM processing.
    
    Args:
        checklist (dict): Checklist data
        
    Returns:
        str: Formatted checklist
    """
    lines = ["# ISO 26262 Part 3 - Item Definition Review Checklist\n"]
    
    # Group items by category
    items_by_category = {}
    for item in checklist.get("items", []):
        category = item.get("category", "Other")
        if category not in items_by_category:
            items_by_category[category] = []
        items_by_category[category].append(item)
    
    # Format by category
    for category, items in items_by_category.items():
        lines.append(f"\n## {category}\n")
        for item in items:
            lines.append(f"**{item['id']}** - {item['requirement']}")
            lines.append(f"  Description: {item['description']}")
            lines.append(f"  Reference: {item.get('iso_clause', 'N/A')}")
            lines.append("")
    
    return "\n".join(lines)


def build_review_prompt(item_definition, checklist):
    """
    Build structured LLM prompt for item definition review.
    
    Args:
        item_definition (str): Item definition content
        checklist (dict): Review checklist
        
    Returns:
        str: Formatted prompt
    """
    checklist_formatted = format_checklist_for_llm(checklist)
    
    # Truncate item definition if too long
    max_length = 12000
    item_def_truncated = item_definition[:max_length]
    if len(item_definition) > max_length:
        item_def_truncated += "\n\n[... content truncated ...]"
    
    prompt = f"""You are a Functional Safety expert conducting an ISO 26262 Part 3 Item Definition review.

# Item Definition Content

{item_def_truncated}

---

# Review Checklist

{checklist_formatted}

---

# Your Task

Review the Item Definition against each checklist item. For each item:

1. **Carefully read** the requirement and description
2. **Search** for evidence in the Item Definition
3. **Assess** whether the requirement is satisfied
4. **Provide specific evidence** - cite sections, headings, or content
5. **Offer actionable improvements** for failed items

# Output Format

For each checklist item, provide your assessment in this exact format:

**ID:** [Checklist ID]
**Category:** [Category Name]
**Requirement:** [Requirement text]
**Description:** [Description from checklist]
**Status:** Pass / Fail / Not Applicable
**Comment:** [Your detailed assessment with specific evidence from the document]
**Hint for improvement:** [Actionable suggestion if Fail, or "N/A" if Pass]

---

# Review Criteria

- **Pass**: Requirement fully met with clear evidence in the document
- **Fail**: Requirement not met, unclear, or insufficient evidence
- **Not Applicable**: Requirement does not apply to this item

# Quality Guidelines

✅ **DO:**
- Quote specific sections or headings as evidence
- Explain WHY something passes or fails
- Provide constructive, actionable improvement hints
- Check for completeness, clarity, and compliance

❌ **DON'T:**
- Give vague assessments without evidence
- Mark as Pass without citing where requirement is met
- Use "Not Applicable" without justification
- Provide generic or unhelpful improvement hints

---

**Begin your review now. Assess ALL checklist items.**
"""
    
    return prompt


# ============================================================================
# TOOLS
# ============================================================================

@tool(
    return_direct=True,
    examples=[
        "review the item definition",
        "check item definition compliance",
        "review item definition with checklist"
    ]
)
def review_item_definition(tool_input, cat):
    """Review Item Definition against ISO 26262 Part 3 compliance checklist.
    Input: not required (loads from item_definition_to_review folder).
    Returns detailed assessment with Pass/Fail status and improvement hints."""
    
    log.info("🔧 TOOL CALLED: review_item_definition")
    
    plugin_folder = os.path.dirname(__file__)
    
    # Step 1: Load checklist
    try:
        checklist = load_checklist(plugin_folder)
        log.info(f"✅ Loaded checklist with {len(checklist.get('items', []))} items")
    except FileNotFoundError:
        return """❌ **Checklist Not Found**

The review checklist file is missing. Please check:
1. File exists: `checklists/item_definition_checklist.json`
2. Plugin is correctly installed
3. File permissions are correct

Contact plugin maintainer if issue persists."""
    except ValueError as e:
        return f"""❌ **Checklist File Corrupted**

The checklist JSON file is invalid: {e}

Please reinstall the plugin or repair the checklist file."""
    
    # Step 2: Load Item Definition
    item_definition = load_item_definition(plugin_folder)
    
    if not item_definition:
        return """❌ **No Item Definition Found**

Please place your Item Definition file in the `item_definition_to_review/` folder.

**Supported formats:**
- `.txt` - Plain text
- `.pdf` - PDF document
- `.docx` - Word document

**Requirements:**
- Only one file should be in the folder
- File must contain the complete Item Definition
- File must be readable (not password-protected)

**Try again after adding your file.**"""
    
    log.info(f"✅ Loaded Item Definition: {len(item_definition)} characters")
    
    # Step 3: Generate review
    try:
        prompt = build_review_prompt(item_definition, checklist)
        log.info("🤖 Generating review with LLM...")
        
        response = cat.llm(prompt)
        log.info(f"✅ Review generated: {len(response)} characters")
        
        # Set working memory for formatter plugin
        cat.working_memory["document_type"] = "item_definition_review"
        cat.working_memory["reviewed_item"] = "Item Definition"
        cat.working_memory["review_date"] = datetime.now().strftime('%Y-%m-%d')
        
        return response
        
    except Exception as e:
        log.error(f"LLM review generation failed: {e}")
        return f"""❌ **Review Generation Failed**

An error occurred while generating the review: {e}

**Possible causes:**
- LLM service unavailable
- Item Definition too large
- Invalid checklist format

**Recommendations:**
1. Try again in a few moments
2. Check if Item Definition is too large (>50 pages)
3. Verify LLM connection in Cheshire Cat settings"""


@tool(
    return_direct=True,
    examples=[
        "generate review template",
        "create item definition review template",
        "make blank review checklist"
    ]
)
def get_review_template(tool_input, cat):
    """Generate blank ISO 26262 Item Definition review template.
    Input: not required.
    Returns template with all checklist items and empty assessment fields."""
    
    log.info("🔧 TOOL CALLED: get_review_template")
    
    plugin_folder = os.path.dirname(__file__)
    
    # Load checklist
    try:
        checklist = load_checklist(plugin_folder)
        log.info(f"✅ Loaded checklist with {len(checklist.get('items', []))} items")
    except FileNotFoundError:
        return """❌ **Checklist Not Found**

Cannot generate template - checklist file is missing.
Please check plugin installation."""
    except ValueError as e:
        return f"""❌ **Checklist File Corrupted**

Cannot generate template - checklist is invalid: {e}"""
    
    # Build template content
    items = checklist.get("items", [])
    
    template_lines = [
        "# ISO 26262 Part 3 - Item Definition Review Template",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*",
        f"*Total Checklist Items: {len(items)}*",
        "",
        "**Instructions:**",
        "- Fill in Status: Pass / Fail / Not Applicable",
        "- Provide detailed Comment with evidence",
        "- Add Hint for improvement for failed items",
        "",
        "---",
        ""
    ]
    
    # Add all checklist items with empty fields
    for item in items:
        template_lines.extend([
            f"**ID:** {item['id']}",
            f"**Category:** {item['category']}",
            f"**Requirement:** {item['requirement']}",
            f"**Description:** {item['description']}",
            f"**ISO Clause:** {item.get('iso_clause', 'N/A')}",
            "**Status:** ",
            "**Comment:** ",
            "**Hint for improvement:** ",
            "",
            "---",
            ""
        ])
    
    template_content = "\n".join(template_lines)
    
    # Set working memory for formatter
    cat.working_memory["document_type"] = "item_definition_review"
    cat.working_memory["reviewed_item"] = "Template - Item Definition Review"
    cat.working_memory["is_template"] = True
    cat.working_memory["review_date"] = datetime.now().strftime('%Y-%m-%d')
    
    log.info(f"✅ Template generated: {len(template_content)} characters")
    
    return template_content


# ============================================================================
# LEGACY COMPATIBILITY (optional - keep old names as aliases)
# ============================================================================

# Uncomment if you want to maintain backward compatibility
# review_item_definition_old = review_item_definition
# generate_review_template = get_review_template