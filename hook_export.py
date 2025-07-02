from cat.mad_hatter.decorators import hook
from docx import Document
from io import BytesIO, StringIO
import base64
import os
import csv
import zipfile
import re
import json
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

# Hook: Export LLM output to .docx, .csv and package into .zip
@hook(priority=0)
def before_cat_sends_message(final_output, cat):
    """
    This hook runs after the LLM generates a response that includes a markdown-style table.
    It performs the following actions:
    
    1. Detects if the response contains a checklist review (based on presence of '|')
    2. Parses the markdown-style table from the LLM output
    3. Creates a structured Word (.docx) document with tables per item
    4. Builds a CSV file using ';' as delimiter for Excel compatibility
    5. Packages both files into a ZIP archive
    6. Returns the ZIP file to the user for download
    7. Saves the ZIP locally for future reference
    
    Args:
        final_output (dict): The final message to be sent to the user.
        cat (Cat): Cheshire Cat instance, used to access memory, settings, etc.
        
    Returns:
        dict: Updated final_output with downloadable ZIP file attached.
    """

    # Only process if the output contains a markdown-style table (i.e., has '|')
    if "|" in final_output.get("content", ""):
        print("📦 Packaging .docx and .csv into .zip...")

        # Get current plugin folder path
        plugin_folder = os.path.dirname(__file__)

        # Load checklist JSON data to enrich review with requirement descriptions and ISO clauses
        checklist = load_checklist(plugin_folder)
        checklist_items = checklist.get("items", [])
        checklist_map = {item["id"]: item for item in checklist_items}

        # Step 1: Parse markdown-style table from LLM output
        review_data = parse_markdown_table(final_output["content"])

        # Step 2: Create a new Word document
        doc = Document()
        doc.add_heading('ISO 26262 Part 3 - Item Definition Review Report', level=1)

        # Loop through each item in the parsed LLM response
        for item in review_data:
            item_id = item.get("ID", "")
            checklist_item = checklist_map.get(item_id, {})
            category = checklist_item.get("category", "Uncategorized")

        # Add category heading and explanation only when it changes
            if not hasattr(before_cat_sends_message, "last_category") or before_cat_sends_message.last_category != category:
                doc.add_heading(category, level=2)
                add_section_explanation(doc, category)  # <<< Insert explanation here
                before_cat_sends_message.last_category = category

            # Create a table with 2 columns and N rows depending on content
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            
            # Set column widths: 1st column = 1.5 inches (narrow), 2nd column = 4.5 inches (wide)
            for row in table.rows:
                row.cells[0].width = Pt(2.5 * 1440 / 72)  # ~1.5 inches wide
                row.cells[1].width = Pt(20.0 * 1440 / 72)  # ~10.5 inches wide
            
            # Set font size and bold headers
            def set_cell_text(cell, text, bold=False):
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(text)
                run.bold = bold
                run.font.size = Pt(10)

            set_cell_text(table.rows[0].cells[0], "Requirement", bold=True)
            set_cell_text(table.rows[0].cells[1], checklist_item.get("requirement", ""))

            set_cell_text(table.rows[1].cells[0], "Description", bold=True)
            set_cell_text(table.rows[1].cells[1], checklist_item.get("description", "N/A"))

            set_cell_text(table.rows[2].cells[0], "ISO Clause", bold=True)
            set_cell_text(table.rows[2].cells[1], checklist_item.get("iso_clause", "N/A"))

            set_cell_text(table.rows[3].cells[0], "Result", bold=True)
            set_cell_text(table.rows[3].cells[1], item.get("Status", "Not Reviewed"))

            set_cell_text(table.rows[4].cells[0], "Comment", bold=True)
            set_cell_text(table.rows[4].cells[1], item.get("Comment", ""))

            # Add paragraph after each table to separate items visually
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)     # No extra space after paragraph
            p.paragraph_format.line_spacing = 1.0      # Single line spacing

        # Step 3: Save .docx to buffer (in-memory)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        doc_bytes = doc_buffer.getvalue()
        doc_buffer.close()

        # Step 4: Build CSV content
        csv_buffer = StringIO()
        csv_writer = csv.writer(csv_buffer, delimiter=";")
        csv_writer.writerow(["ID", "Requirement", "Description", "Clause", "Status", "Comment"])

        # Populate CSV rows using both LLM output and checklist metadata
        for item in review_data:
            item_id = item.get("ID", "")
            req = item.get("Requirement", "")
            status = item.get("Status", "")
            comment = item.get("Comment", "")

            checklist_item = checklist_map.get(item_id, {})
            full_requirement = checklist_item.get("requirement", req)
            iso_clause = checklist_item.get("iso_clause", "N/A")
            description = checklist_item.get("description", "N/A")

            csv_writer.writerow([item_id, full_requirement, description, iso_clause, status, comment])

        csv_content = csv_buffer.getvalue()
        csv_buffer.close()

        # Step 5: Package both files into a ZIP archive
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # Add .docx file to ZIP
            doc_filename = f"item_definition_review_{timestamp}.docx"
            zip_file.writestr(doc_filename, doc_bytes)

            # Add .csv file to ZIP
            csv_filename = f"item_definition_review_{timestamp}.csv"
            zip_file.writestr(csv_filename, csv_content)

        zip_buffer.seek(0)
        encoded_zip = base64.b64encode(zip_buffer.read()).decode("utf-8")
        zip_buffer.close()

        # Step 6: Save ZIP file to disk for local access
        exports_folder = os.path.join(plugin_folder, "exports")
        os.makedirs(exports_folder, exist_ok=True)

        zip_path = os.path.join(exports_folder, f"item_definition_review_{timestamp}.zip")

        with open(zip_path, "wb") as f:
            f.write(base64.b64decode(encoded_zip))

        print(f"💾 ZIP file saved at: {zip_path}")

        # Step 7: Attach ZIP file to chat response for download
        final_output["file"] = {
            "name": f"item_definition_review_{timestamp}.zip",
            "content": encoded_zip,
            "type": "zip"
        }

        # Optional: Update chat message to confirm export
        final_output["content"] = (
            "✅ Review completed and exported.\n"
            "See attached `.zip` file containing:\n"
            "- `item_definition_review.docx`\n"
            "- `item_definition_review.csv`"
        )

    return final_output

# Helper Function: Parse markdown-style tables from LLM output
def parse_markdown_table(text):
    """
    Parses markdown-style tables from input text and returns them as a list of dictionaries.

    Args:
        text (str): The raw input string that may contain one or more markdown tables.

    Returns:
        List[dict]: A list of dictionaries, where each dictionary represents a row in the table,
                    with keys from the header and values from the corresponding cells.
                    
    """

    # Regular expression pattern to detect markdown-style tables
    # Matches lines starting and ending with pipes (`|`) and includes:
    # - Header row
    # - Separator row (e.g., |---|---|)
    # - Data rows
    # Uses re.DOTALL so that '.' matches newlines too
    table_pattern = re.compile(r"(\|.*\|\s*\|[-:]*[-|]\s*(?:\|.*\|[\s\d]*)+)", re.DOTALL)

    # Find all matching tables in the input text
    tables = table_pattern.findall(text)

    # Initialize an empty list to store parsed table data
    parsed_data = []

    # Loop through each matched table block
    for table in tables:
        # Split the table into lines and strip whitespace from each line
        lines = [line.strip() for line in table.strip().split('\n')]

        # Skip if there are not enough lines (header + separator + at least one data row)
        if len(lines) < 2:
            continue

        # Extract headers from the first line
        # Split by pipe (`|`) and remove the first and last empty elements using [1:-1]
        headers = [h.strip() for h in lines[0].split('|')[1:-1]]

        # Process each data row starting after the separator line (lines[2:])
        for line in lines[2:]:
            # Split the line by pipe (`|`) and remove the first and last empty elements
            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Skip empty rows
            if not any(cells):
                continue

            # Map headers to cell values and add to result list
            row = dict(zip(headers, cells))
            parsed_data.append(row)

    # Return the list of parsed rows
    return parsed_data

def load_checklist(plugin_folder):
    """
    Loads the ISO 26262 Part 3 compliance checklist from a JSON file.

    Args:
        plugin_folder (str): Path to the plugin folder.

    Returns:
        dict: Checklist data as a Python dictionary.

    Raises:
        FileNotFoundError: If the JSON file is missing.
        json.JSONDecodeError: If the JSON file has syntax errors.
    """
    checklist_path = os.path.join(plugin_folder, "checklists", "item_definition_checklist.json")
    
    try:
        with open(checklist_path, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        raise FileNotFoundError(f"Checklist file not found at {checklist_path}")
    except json.JSONDecodeError as e:
        raise json.JSONDecodeError(f"Invalid JSON in checklist file: {e}", doc=e.doc, pos=e.pos)
    
# Helper Function: Add explanation text for each category
def add_section_explanation(doc, category):
    """
    Adds a brief explanation paragraph after each category heading in the Word document.
    """
    explanations = {
        "Identification and Classification": (
            "This section ensures that the item is uniquely identified within the system architecture or documentation "
            "and properly classified (e.g., as hardware, software, or a system function). This supports traceability from "
            "high-level safety goals down to detailed design elements. Clear identification also enables effective configuration "
            "management and version control throughout the development lifecycle. It is a foundational element for ensuring structured functional safety development."
        ),
        "Functional Description": (
            "This section describes the expected behavior of the item under all operating conditions, including normal, degraded, "
            "and fault modes. It includes definitions of interfaces, timing constraints, and performance requirements. A well-defined "
            "functional description is essential for identifying potential failure scenarios and serves as input to hazard analysis. It helps "
            "ensure that all relevant behaviors are considered when deriving safety requirements."
        ),
        "Safety-Related Attributes": (
            "This section captures key safety-related properties such as safety goals, mitigation strategies, diagnostic coverage, "
            "and safe state definitions. These attributes are derived from the Hazard Analysis and Risk Assessment (HARA) and form the basis "
            "of the functional safety concept. They guide the implementation of safety mechanisms and define how the item contributes to overall "
            "system safety. Proper documentation ensures alignment with ISO 26262 expectations for safety integrity."
        ),
        "Dependencies and Interactions": (
            "This section identifies internal and external dependencies, including interactions with other systems, environmental influences, "
            "and user inputs. Understanding these relationships is critical for defining correct assumptions and boundary conditions during development. "
            "It also supports the identification of potential interference or integration risks that could impact safety. Accurate documentation ensures robust "
            "interface management and system integration."
        ),
        "System Boundaries and Context": (
            "This section defines the physical and logical boundaries of the item, along with environmental conditions and design constraints. "
            "It clarifies where the item operates and under what limitations, such as temperature, vibration, or EMC exposure. These details ensure that "
            "the item is developed and validated under realistic assumptions. Defining this context early supports the creation of accurate test plans and operational profiles."
        ),
        "Review and Approval": (
            "This section confirms that a formal review process was followed and that all necessary approvals were obtained before finalizing the item definition. "
            "It verifies that review minutes, action items, and change records are documented and closed. Configuration management practices should also be applied to maintain "
            "document integrity. This ensures process compliance and provides an auditable trail for quality assurance and functional safety governance."
        )
    }

    explanation = explanations.get(category)
    if explanation:
        paragraph = doc.add_paragraph(explanation)
        paragraph.style = 'Normal'
        paragraph.alignment = 0  # Left-aligned
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.2